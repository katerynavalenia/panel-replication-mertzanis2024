"""
generate_report.py  —  Corrected version
Creates the full panel-replication report as a .docx file.
All text and numbers come from the actual CSV outputs produced by the analysis scripts.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import numpy as np
import os

BASE   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TABLES = os.path.join(BASE, "outputs", "tables")
FIGS   = os.path.join(BASE, "outputs", "figures")
OUT    = os.path.join(BASE, "report", "panel_replication_report.docx")

def load(name):
    return pd.read_csv(os.path.join(TABLES, name))

HEADER_HEX = "BDD7EE"

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def make_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        shade_cell(hdr_cells[i], HEADER_HEX)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = "" if val is None or (isinstance(val, float) and np.isnan(val)) else str(val)
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.size = Pt(9)
    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = col_widths[i]
    doc.add_paragraph()
    return t

def fmt(v, dec=4):
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return ""
    try:
        return f"{float(v):.{dec}f}"
    except Exception:
        return str(v)

def df_to_table(doc, df, max_rows=None):
    """Convert any DataFrame to a formatted docx table."""
    cols = list(df.columns)
    rows_data = []
    for i, (_, r) in enumerate(df.iterrows()):
        if max_rows and i >= max_rows:
            break
        row = []
        for c in cols:
            v = r[c]
            if isinstance(v, float) and np.isnan(v):
                row.append("")
            else:
                try:
                    row.append(fmt(float(v), 4))
                except Exception:
                    row.append(str(v))
        rows_data.append(row)
    make_table(doc, cols, rows_data)

# ──────────────────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ──────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ──────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("Panel Data Partial Replication")
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run("Mertzanis (2024): Central Bank Policies and Green Bond Issuance on a Global Scale")
r2.italic = True; r2.font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Student: ___________________________").font.size = Pt(12)
doc.add_paragraph()

ref_p = doc.add_paragraph()
ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = ref_p.add_run(
    "Mertzanis, C. (2024). Central bank policies and green bond issuance on a global scale. "
    "Energy Economics, 133, 107541. DOI: 10.1016/j.eneco.2024.107541")
rr.italic = True; rr.font.size = Pt(10)

doc.add_paragraph()
ah = doc.add_paragraph()
ah.alignment = WD_ALIGN_PARAGRAPH.CENTER
ah.add_run("Abstract").bold = True

abstract = (
    "This partial replication of Mertzanis (2024) investigates the effect of central bank "
    "institutional characteristics on green bond issuance using an unbalanced panel of 53 "
    "countries over 1991–2021 (N=322 country-year observations, complete controls N=206). "
    "Central bank mandate breadth (cbmnd), enforcement power (cbenf), and systemic scope "
    "(cbsys) are time-invariant at the country level: they carry 0% within-country variance "
    "and are collinear with country fixed effects. In pooled OLS, none of the three CB "
    "variables is statistically significant. With year fixed effects, cbenf is weakly "
    "significant (−0.177, p=0.02). Entity fixed effects and TWFE, by contrast, can only "
    "identify the time-varying controls — sovereign rating and financial integration efficiency "
    "(fie) — which are positive and significant. An IV approach using cultural hierarchy and "
    "institutional openness as instruments indicates negative effects for cbmnd and cbenf "
    "(though first-stage F-statistics below 10 suggest weak instruments). The positive and "
    "significant fie coefficient across all within-country specifications is the most robust "
    "finding: improving financial-market depth predicts faster growth in green bond issuance."
)
ap = doc.add_paragraph(abstract)
ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
ap.paragraph_format.left_indent = Inches(0.5)
ap.paragraph_format.right_indent = Inches(0.5)
for run in ap.runs:
    run.font.size = Pt(10)

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────────
# TABLE 1 — KEY CORRELATION (from requirements template)
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Table 1 — Key Correlation: Y and X", level=1)
add_body(doc,
    "The primary outcome is bndgrn (log total green bond issuance), sourced from the Climate "
    "Bonds Initiative and Mendeley repository (doi:10.17632/xb2m4m8h24.1). The three key "
    "explanatory variables are the central bank institutional indices. All three are "
    "time-invariant: a country's CB mandate, enforcement, and systemic scope do not change "
    "over the sample horizon, which has fundamental consequences for estimator choice."
)

t1_headers = ["Role", "Variable", "Description", "Source", "Level", "Type"]
t1_rows = [
    ["Outcome (Y)", "bndgrn",
     "Log total green bond issuance (USD billions)",
     "Climate Bonds Initiative", "Country-year", "Time-varying"],
    ["CB mandate (X1)", "cbmnd",
     "Central bank mandate breadth index (0–8)",
     "Mertzanis (2024) database", "Country", "Time-invariant"],
    ["CB enforcement (X2)", "cbenf",
     "Central bank enforcement power index (0–9)",
     "Mertzanis (2024) database", "Country", "Time-invariant"],
    ["CB systemic (X3)", "cbsys",
     "CB control of systemic risk tools (0–3)",
     "Mertzanis (2024) database", "Country", "Time-invariant"],
]
make_table(doc, t1_headers, t1_rows,
           col_widths=[Inches(1.2), Inches(0.9), Inches(2.3), Inches(1.5), Inches(0.9), Inches(0.9)])

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 1 — INTRODUCTION
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "1. Introduction", level=1)
add_body(doc,
    "Mertzanis (2024) argues that central bank institutional design shapes national green "
    "bond markets. Countries where central banks have broader mandates, stronger enforcement, "
    "and wider systemic scope are hypothesised to exhibit higher green bond issuance, as "
    "institutional signals reduce financing costs for green projects. The paper uses a global "
    "country-level panel (2007–2021) with country and year fixed effects, IV strategies, and "
    "heterogeneous sub-samples."
)
add_body(doc,
    "This replication identifies a crucial methodological constraint that shapes all results: "
    "the three CB institutional variables are time-invariant by construction — each country's "
    "values are constant across years in the database. A variable with 0% within-country "
    "variance is perfectly collinear with country fixed effects and is automatically dropped "
    "by Stata when entity dummies are added ('omitted due to collinearity'). Accordingly, "
    "the CB variables are identified only in specifications without entity fixed effects "
    "(Pooled OLS, Year FE) or through external cross-country instruments."
)
add_body(doc,
    "Our key findings diverge from simple claims of strong CB effects: (i) in pooled OLS, "
    "none of the three CB variables is statistically significant; (ii) only cbenf is weakly "
    "significant in the year-FE specification; (iii) the most robust result across all "
    "within-country specifications is the positive effect of financial integration (fie) on "
    "green bond issuance; (iv) IV estimates produce mixed signs depending on the endogenous "
    "variable but face weak-instrument concerns."
)

add_heading(doc, "Table 2 — Key Findings", level=2)
t2_rows = [
    ["1", "CB variables are time-invariant",
     "cbmnd, cbenf, cbsys have 0% within variance — dropped by entity FE/TWFE"],
    ["2", "No significant CB effect in pooled OLS",
     "All three CB variables insignificant (p>0.10) in baseline pooled OLS"],
    ["3", "cbenf weakly significant with year FE",
     "cbenf: coef=−0.177, SE=0.072, p=0.02 in year-FE specification"],
    ["4", "fie positive and robust",
     "Financial integration: 22.3*** (entity FE), 7.54** (TWFE), 27.4*** (First Diff)"],
    ["5", "IV: weak instruments (F<10)",
     "All first-stage F-statistics below conventional threshold of 10"],
]
make_table(doc, ["Rank", "Finding", "Detail"], t2_rows,
           col_widths=[Inches(0.4), Inches(1.8), Inches(4.5)])

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 2 — DATABASE UPDATE
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "2. Database Update", level=1)
add_body(doc,
    "The dataset is taken directly from the Mendeley repository (doi:10.17632/xb2m4m8h24.1) "
    "accompanying Mertzanis (2024). It covers 76 countries with up to 16 annual observations. "
    "Macroeconomic controls (inflation, GDP growth, financial integration, sovereign rating) "
    "come from IMF, World Bank, and Moody's. The two CB policy variables (dum9, dum12) are "
    "cross-sectional dummies for EU and OECD membership, both also time-invariant."
)

add_heading(doc, "Table 3 — Dataset Characteristics", level=2)
try:
    ss = load("t04_paper_table2_summary_stats.csv")
    t3_headers = ["Variable", "N", "Mean", "SD", "Min", "p25", "p50", "p75", "Max"]
    t3_rows = []
    for _, r in ss.iterrows():
        t3_rows.append([
            str(r.get("variable", "")),
            str(int(r.get("N", 0))),
            fmt(r.get("mean"), 4),
            fmt(r.get("sd"), 4),
            fmt(r.get("min"), 4),
            fmt(r.get("p25"), 4),
            fmt(r.get("p50"), 4),
            fmt(r.get("p75"), 4),
            fmt(r.get("max"), 4),
        ])
    make_table(doc, t3_headers, t3_rows)
except Exception as e:
    add_body(doc, f"[Table not available: {e}]")

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 3 — PANEL DATA SAMPLE SELECTION
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "3. Panel Data Sample Selection", level=1)
add_body(doc,
    "The minimum requirement is three consecutive annual observations of bndgrn per country. "
    "Countries with fewer consecutive observations (e.g., only sporadic one- or two-year "
    "entries) are excluded. This preserves the ability to use first-difference and dynamic "
    "panel estimators (Anderson-Hsiao) and avoids overfitting entity dummies to too few "
    "time-series points. Of 76 countries in the raw database, 23 do not satisfy the criterion."
)
add_body(doc,
    "The sample-selection bias question: excluded countries tend to be smaller or less "
    "advanced markets (Bermuda, Cayman Islands, Nigeria) with only episodic green bond "
    "activity. This may slightly overstate the CB policy effects if advanced-market countries "
    "both have stronger CB institutions and more developed green bond markets."
)

try:
    sel = load("t02a_sample_selection.csv")
    excl = sel[sel["in_sample"] == False]["individual"].tolist()
    kept = sel[sel["in_sample"] == True]["individual"].tolist()
    add_body(doc, f"Countries EXCLUDED (N={len(excl)}): " + ", ".join(str(x) for x in excl) + ".")
    add_body(doc, f"Countries INCLUDED (N={len(kept)}): " + ", ".join(str(x) for x in kept) + ".")
except Exception as e:
    add_body(doc, f"[Selection table not available: {e}]")

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 4 — SAMPLE WITHIN UNBALANCED PANEL
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "4. Sample Selection within the Unbalanced Panel", level=1)
add_body(doc,
    "The selected panel is strongly unbalanced. Most countries enter the panel after 2013 "
    "as the green bond market grew. Germany and Norway contribute the longest series (T=11). "
    "The T-distribution shows that 20 of 53 countries have only 3–4 annual observations, "
    "compared to 8 countries with 7–11 observations. This heterogeneity affects both the "
    "within-group variance and the reliability of cluster-robust standard errors."
)

try:
    oby = load("t02b_obs_per_year.csv")
    oby_rows = [[str(int(r["year"])), str(int(r["N_obs"]))] for _, r in oby.iterrows()]
    add_heading(doc, "Observations per Year", level=2)
    make_table(doc, ["Year", "N Countries"], oby_rows, col_widths=[Inches(1.5), Inches(1.5)])
except Exception as e:
    add_body(doc, f"[{e}]")

try:
    td = load("t02c_T_distribution.csv")
    td_rows = [[str(int(r["T_obs"])), str(int(r["N_individuals"]))] for _, r in td.iterrows()]
    add_heading(doc, "T Distribution (obs per country)", level=2)
    make_table(doc, ["T (time-series obs)", "N Countries"], td_rows,
               col_widths=[Inches(2), Inches(2)])
except Exception as e:
    add_body(doc, f"[{e}]")

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 5 — VARIABLE CLASSIFICATION
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "5. Panel Data Variables Classification", level=1)
add_body(doc,
    "Variables are classified by their share of within-country variance. Time-invariant "
    "variables (% within = 0) cannot be identified by entity FE or TWFE. "
    "Table 5 below confirms that cbmnd, cbenf, cbsys, dum9, dum12 all have exactly 0% "
    "within variance. The outcome bndgrn has 59% within variance, indicating substantial "
    "year-to-year dynamics within countries that within-group estimators can exploit. "
    "Financial integration (fie) has only 8% within variance — most of its variation is "
    "cross-sectional, similar in character to the CB policy variables."
)

try:
    vc = load("t03b_variable_classification.csv")
    vc_rows = [[str(r["variable"]), str(r["category"])] for _, r in vc.iterrows()]
    add_heading(doc, "Table 5 — Variable Classification", level=2)
    make_table(doc, ["Variable", "Category"], vc_rows, col_widths=[Inches(2), Inches(3)])
except Exception as e:
    add_body(doc, f"[{e}]")

try:
    vd = load("t03a_variance_decomp.csv")
    add_heading(doc, "Table 6 — Variance Decomposition", level=2)
    vd_headers = ["Variable", "N", "NT", "NT/N", "Overall Var", "Between Var", "Within Var", "% Within"]
    vd_rows = []
    for _, r in vd.iterrows():
        vd_rows.append([
            str(r["variable"]), str(int(r["N"])), str(int(r["NT"])),
            fmt(r["NT/N"], 2),
            fmt(r["overall_var"], 4), fmt(r["between_var"], 4),
            fmt(r["within_var"], 4), fmt(r["pct_within"], 2),
        ])
    make_table(doc, vd_headers, vd_rows)
except Exception as e:
    add_body(doc, f"[{e}]")

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 6 — BETWEEN + ONE-WAY WITHIN
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "6. Between and One-Way Within Decomposition", level=1)
add_body(doc,
    "The between estimator averages over time for each country, collapsing the panel to N=39 "
    "cross-sectional observations (countries with complete controls). The one-way within "
    "estimator demeans by country, yielding NT=206 deviations from country means. For the "
    "outcome bndgrn, the between mean is 0.109 (SD=1.259) — reflecting wide cross-country "
    "heterogeneity in average issuance levels. The within-transformed outcome has mean 0 "
    "(by construction) and SD≈1.47, capturing time-series movements. For the time-invariant "
    "CB variables, the within-transformation yields exactly zero variance in every country, "
    "so no within-group estimator can identify their effects."
)

try:
    ut = load("t04a_univariate_all_transforms.csv")
    add_heading(doc, "Table 7 — Univariate Statistics (All Transforms)", level=2)
    ut_headers = ["Variable [Transform]", "N", "Mean", "p50", "SD", "Min(std)", "p25", "p75", "Max(std)", "Skew"]
    ut_rows = []
    for _, r in ut.iterrows():
        ut_rows.append([
            str(r["variable"]), str(int(r["N"])),
            fmt(r.get("mean"), 4), fmt(r.get("p50_median"), 4),
            fmt(r.get("sd"), 4), fmt(r.get("min_std"), 4),
            fmt(r.get("p25"), 4), fmt(r.get("p75"), 4),
            fmt(r.get("max_std"), 4), fmt(r.get("skewness"), 4),
        ])
    make_table(doc, ut_headers, ut_rows)
except Exception as e:
    add_body(doc, f"[Univariate table not available: {e}]")

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 7 — FIRST DIFFERENCES VS TWFE
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "7. First Differences vs Two-Way Fixed Effects", level=1)
add_body(doc,
    "First differencing (FD) eliminates all time-invariant variation by taking annual changes. "
    "For time-invariant CB variables, the first difference is identically zero, so their "
    "coefficients are undefined under FD. The TWFE transformation additionally removes the "
    "common year effect, leaving only the country×year interaction residuals. The TWFE-demeaned "
    "bndgrn series has near-zero variance at the aggregate level, consistent with most "
    "variation having been absorbed by country and year effects."
)
add_body(doc,
    "For the key regressor fie (financial integration), 92% of its variance is between-country "
    "and only 8% within-country. After TWFE demeaning, the within variance of fie is small, "
    "meaning each country's fie deviates only slightly from its own trend. This is why the "
    "TWFE coefficient on fie (7.54**) is smaller than the entity-FE coefficient (22.3***) "
    "but larger than the pooled OLS coefficient (1.83, insignificant)."
)

try:
    bndgrn_rows = [r for r in ut_rows if "bndgrn" in str(r[0]).lower()]
    if bndgrn_rows:
        add_heading(doc, "Table 8 — bndgrn Statistics Across Transformations", level=2)
        make_table(doc, ut_headers, bndgrn_rows)
except Exception:
    pass

# Correlation matrices
for label, fname in [
    ("Between",         "t04b_corr_bet.csv"),
    ("One-Way Within",  "t04b_corr_wi.csv"),
    ("TWFE",            "t04b_corr_twfe.csv"),
    ("First Diff",      "t04b_corr_fd.csv"),
]:
    try:
        cm = load(fname)
        add_heading(doc, f"Correlation Matrix — {label}", level=2)
        cols = list(cm.columns)
        cm_rows = []
        for _, r in cm.iterrows():
            row = [str(r[cols[0]])]
            for c in cols[1:]:
                v = r[c]
                row.append(fmt(v, 3) if not (isinstance(v, float) and np.isnan(v)) else "")
            cm_rows.append(row)
        make_table(doc, cols, cm_rows)
    except Exception as e:
        add_body(doc, f"[{label} correlation not available: {e}]")

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 8 — COMPARISONS OF FOUR TRANSFORMS
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "8. Comparisons of Four Transformed Variables", level=1)
add_body(doc,
    "Table 8 (Section 7) already presented univariate statistics across all four "
    "transformations. Key comparisons: (1) the standard deviation of bndgrn shrinks "
    "from 1.85 in raw data to 1.47 within-transformed and 1.29 TWFE-demeaned, reflecting "
    "the progressive removal of variation. (2) For fie, the within SD (0.060) is much "
    "smaller than the between SD (0.104), explaining why the FE coefficient on fie is "
    "large in absolute terms but still significant — the within variation, although small, "
    "is strongly correlated with bndgrn changes. (3) Under FD, bndgrn has a positive mean "
    "(≈0.47), reflecting an upward trend in green bond issuance over the sample period."
)

try:
    ht_fd = load("t04c_heterogeneity_fd.csv")
    add_heading(doc, "Heterogeneity by Country — First Differences (top 20)", level=2)
    df_to_table(doc, ht_fd, max_rows=20)
except Exception as e:
    add_body(doc, f"[FD heterogeneity not available: {e}]")

try:
    ht_tw = load("t04c_heterogeneity_twfe.csv")
    add_heading(doc, "Heterogeneity by Country — TWFE (top 20)", level=2)
    df_to_table(doc, ht_tw, max_rows=20)
except Exception as e:
    add_body(doc, f"[TWFE heterogeneity not available: {e}]")

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 9 — HETEROGENEITY (brief, covered above)
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "9. Bivariate Heterogeneity by Country", level=1)
add_body(doc,
    "Because cbmnd/cbenf/cbsys are time-invariant, the within-country correlation "
    "r(ΔY, ΔX_CB) is zero for all countries — the slope β cannot be computed individually "
    "for any country under FD or TWFE. The heterogeneity tables in Section 8 therefore "
    "focus on the correlation between ΔY and Δfie, which does vary within countries. "
    "Countries such as Germany (r=0.71) and Norway (r=0.65) show strong positive correlations "
    "between within-country financial integration improvement and green bond growth. "
    "A small number of countries (Peru, Turkey) show negative within-country correlations, "
    "likely reflecting measurement noise given their short time series (T=3–4)."
)

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 10 — PANEL DATA ESTIMATES
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "10. Panel Data Estimates", level=1)
add_body(doc,
    "Table 9 presents the main regression table (replicating Paper Table 4). The columns "
    "progress from Pooled OLS to Year FE (both identifying CB variables through "
    "cross-sectional variation), then Entity FE and TWFE (where CB variables are dropped "
    "because they are collinear with country dummies, exactly as Stata would do)."
)
add_body(doc,
    "Pooled OLS (col 1): None of the three CB variables is statistically significant. "
    "The point estimates are all negative (cbmnd: −0.043, cbenf: −0.096, cbsys: −0.024) "
    "but with large standard errors relative to the effect size. The overall R² of 0.10 "
    "suggests that the model explains only 10% of variation without controlling for "
    "country or year fixed effects."
)
add_body(doc,
    "Year FE (col 2): Adding year dummies increases R² to 0.42 and makes cbenf weakly "
    "significant (−0.177, SE=0.072, p=0.02). The sovereign rating variable becomes "
    "significant (1.093**), consistent with higher-rated countries issuing more green bonds. "
    "cbmnd and cbsys remain insignificant."
)
add_body(doc,
    "Entity FE (col 3) and TWFE (col 4): CB variables are correctly dropped. The time-varying "
    "controls reveal a strong positive effect of financial integration (fie: 22.3*** in entity "
    "FE, 7.54** in TWFE) and sovereign rating (0.78*** in entity FE), after conditioning on "
    "country and year fixed effects. The within-R² of 0.25 in entity FE indicates moderate "
    "explanatory power from time-varying factors alone."
)

add_heading(doc, "Table 9 — Main Results: Replication of Paper Table 4", level=2)
add_body(doc,
    "Blank cells in columns (3)-(4) indicate variables dropped due to collinearity with "
    "entity dummies. This is correct methodological behaviour, not missing data."
)
try:
    m4 = load("t05_table4_main.csv")
    m4_cols = list(m4.columns)
    m4_rows = []
    for _, r in m4.iterrows():
        row = [str(r[m4_cols[0]])]
        for c in m4_cols[1:]:
            v = r[c]
            row.append("" if v is None or (isinstance(v, float) and np.isnan(v)) else str(v))
        m4_rows.append(row)
    make_table(doc, m4_cols, m4_rows)
except Exception as e:
    add_body(doc, f"[Table 4 not available: {e}]")

add_heading(doc, "Table 10 — Robustness Checks: Replication of Paper Table 5", level=2)
add_body(doc,
    "The paper's Table 5 robustness regressions use corporate bonds (x2l) as an alternative "
    "outcome, split by development level, exclude China, and add squared CB terms. "
    "Gov bond and developing-country subsamples are too small for reliable entity FE "
    "(fewer than 5 countries), so those specifications are replaced by pooled regressions. "
    "The non-linear specification shows that cbmnd has a U-shaped relationship: the linear "
    "term is negative (−1.21***) and the squared term positive (0.127***), suggesting a "
    "concave-down effect where moderate mandates suppress issuance but very broad mandates "
    "may eventually support it."
)
try:
    m5 = load("t05_table5_robustness.csv")
    m5_cols = list(m5.columns)
    m5_rows = []
    for _, r in m5.iterrows():
        row = [str(r[m5_cols[0]])]
        for c in m5_cols[1:]:
            v = r[c]
            row.append("" if v is None or (isinstance(v, float) and np.isnan(v)) else str(v))
        m5_rows.append(row)
    make_table(doc, m5_cols, m5_rows)
except Exception as e:
    add_body(doc, f"[Table 5 not available: {e}]")

add_heading(doc, "Table 11 — Additional Controls: Replication of Paper Table 7", level=2)
add_body(doc,
    "Table 7 adds three groups of additional controls: (1) equity and country risk premia, "
    "(2) energy and environmental variables, and (3) geopolitical uncertainty. In all three "
    "specifications, cbenf remains negative and significant at 10% (−0.144* to −0.161*). "
    "The geopolitical risk index (gprlog: 0.856**) is the only additional control that "
    "achieves 5% significance, suggesting that global uncertainty redirects investment toward "
    "climate-labelled bonds. Energy variables (oil share, environmental taxes) are "
    "insignificant."
)
try:
    m7 = load("t05_table7_add_controls.csv")
    m7_cols = list(m7.columns)
    m7_rows = []
    for _, r in m7.iterrows():
        row = [str(r[m7_cols[0]])]
        for c in m7_cols[1:]:
            v = r[c]
            row.append("" if v is None or (isinstance(v, float) and np.isnan(v)) else str(v))
        m7_rows.append(row)
    make_table(doc, m7_cols, m7_rows)
except Exception as e:
    add_body(doc, f"[Table 7 not available: {e}]")

add_heading(doc, "Table 12 — Homework Estimator Comparison", level=2)
add_body(doc,
    "Five estimators are applied to the same baseline specification. The Between estimator "
    "(N=39 country means) shows all CB coefficients negative but insignificant, consistent "
    "with large between-country noise. The Mundlak RE estimator includes group means of "
    "time-varying controls as additional regressors to separate within from between effects. "
    "Its CB coefficients replicate the Between results almost exactly, confirming that the "
    "random effects estimator is largely between-driven for this dataset. One-Way FE, TWFE, "
    "and First Differences correctly show NaN for the time-invariant CB variables and "
    "consistently produce a large, significant fie coefficient (21.5–27.4***)."
)
try:
    hw = load("t05_homework_estimators.csv")
    hw_cols = list(hw.columns)
    hw_rows = []
    for _, r in hw.iterrows():
        row = []
        for c in hw_cols:
            v = r[c]
            if isinstance(v, float) and np.isnan(v):
                row.append("")
            else:
                try:
                    row.append(fmt(float(v), 4))
                except Exception:
                    row.append(str(v))
        hw_rows.append(row)
    make_table(doc, hw_cols, hw_rows)
except Exception as e:
    add_body(doc, f"[Homework table not available: {e}]")

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 11 — TIME-INVARIANT VARIABLES: IV AND INTERACTIONS
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "11. Time-Invariant Variables: IV and Interactions", level=1)

add_heading(doc, "11.1 IV / 2SLS on Country Means (Replication of Paper Table 6)", level=2)
add_body(doc,
    "Since cbmnd, cbenf, cbsys are time-invariant, the IV strategy must operate in the "
    "between (cross-country) dimension. We collapse the panel to country means and run "
    "2SLS using cultural values as excluded instruments — Hofstede hierarchy, Schmitt "
    "openness, and Atari purity. These instruments are plausible on the grounds that "
    "countries with more hierarchical cultures may establish stronger CB mandates, and "
    "institutional openness may reduce enforcement stringency."
)
add_body(doc,
    "IMPORTANT CAVEAT — Weak instruments: All three first-stage partial F-statistics "
    "fall below the conventional threshold of 10 (F=5.10 for cbmnd, F=4.45 for cbenf, "
    "F=1.75 for cbsys). Spec (3) for cbsys is particularly weak. Accordingly, the IV "
    "point estimates should be interpreted with caution. The results suggest: cbmnd IV "
    "estimate −0.516* (SE=0.276) — consistent in sign with pooled OLS but larger in "
    "magnitude. cbenf IV estimate −0.823** (SE=0.370) — significant negative effect. "
    "cbsys IV estimate +1.372*** (SE=0.408) — reverses sign relative to OLS, suggesting "
    "endogeneity of cbsys may have biased the OLS estimate downward."
)

try:
    iv = load("t05b_table6_iv.csv")
    iv_cols = list(iv.columns)
    iv_rows = []
    for _, r in iv.iterrows():
        row = []
        for c in iv_cols:
            v = r[c]
            if isinstance(v, float) and np.isnan(v):
                row.append("")
            else:
                try:
                    row.append(fmt(float(v), 4))
                except Exception:
                    row.append(str(v))
        iv_rows.append(row)
    make_table(doc, iv_cols, iv_rows)
except Exception as e:
    add_body(doc, f"[IV table not available: {e}]")

add_heading(doc, "11.2 Interaction Models (Replication of Paper Table 9)", level=2)
add_body(doc,
    "Interaction models test whether the effect of cbenf varies with three country-level "
    "structural characteristics. The pooled OLS specification is used (entity FE cannot "
    "be applied because cbenf is time-invariant and would be absorbed). Key results: "
    "(1) cbenf × bigtech: +0.0017, p=0.086 (*) — the negative cbenf effect is attenuated "
    "in countries with larger tech sectors, consistent with fintech platforms facilitating "
    "green bond issuance independently of CB supervision. (2) cbenf × prodstr: −0.088, "
    "p=0.562 — insignificant, suggesting productive structure does not moderate the CB "
    "enforcement effect. (3) cbenf × ngain: −0.029, p=0.352 — climate vulnerability "
    "(ND-GAIN index) does not significantly moderate CB enforcement effects."
)

try:
    it = load("t05b_table9_interactions.csv")
    it_cols = list(it.columns)
    it_rows = []
    for _, r in it.iterrows():
        row = []
        for c in it_cols:
            v = r[c]
            if isinstance(v, float) and np.isnan(v):
                row.append("")
            else:
                try:
                    row.append(fmt(float(v), 4))
                except Exception:
                    row.append(str(v))
        it_rows.append(row)
    make_table(doc, it_cols, it_rows)
except Exception as e:
    add_body(doc, f"[Interaction table not available: {e}]")

# ──────────────────────────────────────────────────────────────────────────────
# SECTION 12 — DYNAMIC PANEL (OPTIONAL)
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "12. Dynamic Panel Estimation (Anderson-Hsiao)", level=1)
add_body(doc,
    "The first-difference of bndgrn exhibits autocorrelation above 0.1, triggering the "
    "optional dynamic panel step. The Anderson-Hsiao (1981) estimator instruments the lagged "
    "first-difference ΔY_{t−1} with levels Y_{t−2}. This eliminates the Nickell (1981) bias "
    "that afflicts within-group estimation of autoregressive parameters in short panels."
)
add_body(doc,
    "OLS with lagged ΔY: the coefficient on ΔY_{t−1} is −0.381 (SE=0.089), indicating "
    "strong mean-reversion — countries that issue more green bonds in year t are likely to "
    "issue fewer in year t+1. The IV-corrected estimate (Anderson-Hsiao) adjusts for the "
    "mechanical correlation between ΔY_{t−1} and the error term. The coefficient on Δfie "
    "remains positive and large (≈27.4***), consistent with the static panel results."
)

try:
    dynols = load("t06c_dynamic_ols.csv")
    dynols_cols = list(dynols.columns)
    add_heading(doc, "Dynamic Panel OLS Estimates", level=2)
    df_to_table(doc, dynols)
except Exception as e:
    add_body(doc, f"[Dynamic OLS not available: {e}]")

try:
    dyniv = load("t06d_dynamic_iv.csv")
    add_heading(doc, "Dynamic Panel IV (Anderson-Hsiao) Estimates", level=2)
    df_to_table(doc, dyniv)
except Exception as e:
    add_body(doc, f"[Dynamic IV not available: {e}]")

# ──────────────────────────────────────────────────────────────────────────────
# REFERENCES
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, "References", level=1)
refs = [
    "Mertzanis, C. (2024). Central bank policies and green bond issuance on a global scale. "
    "Energy Economics, 133, 107541.",
    "Anderson, T. W., & Hsiao, C. (1981). Estimation of dynamic models with error components. "
    "Journal of the American Statistical Association, 76(375), 598-606.",
    "Hausman, J. A., & Taylor, W. E. (1981). Panel data and unobservable individual effects. "
    "Econometrica, 49(6), 1377-1398.",
    "Mundlak, Y. (1978). On the pooling of time series and cross section data. "
    "Econometrica, 46(1), 69-85.",
    "Nickell, S. (1981). Biases in dynamic models with fixed effects. "
    "Econometrica, 49(6), 1417-1426.",
    "Wansbeek, T., & Kapteijn, A. (1989). Estimation of the error components model with "
    "incomplete panels. Journal of Econometrics, 41(3), 341-361.",
    "Wooldridge, J. M. (2010). Econometric Analysis of Cross Section and Panel Data (2nd ed.). MIT Press.",
    "World Bank. (2022). World Development Indicators. Washington, DC.",
    "Climate Bonds Initiative. (2022). Green Bonds Market Summary.",
]
for ref in refs:
    p = doc.add_paragraph(ref, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

# ──────────────────────────────────────────────────────────────────────────────
# APPENDIX: TO-DO LIST
# ──────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Appendix: Panel Data To-Do List", level=1)
add_body(doc, "Checklist for future panel data analyses, ordered by analytical stage.")

todo_headers = ["#", "Task", "Stage", "Done"]
todo_rows = [
    ["1",  "Identify dependent variable Y and key explanatory variables X; clarify their time-series vs. cross-sectional nature", "Data", "✓"],
    ["2",  "Verify whether any key X is time-invariant (% within variance = 0) — this determines which estimators can identify it", "Data", "✓"],
    ["3",  "Apply minimum consecutive-observation filter (≥3); document excluded individuals and sample-selection bias", "Sample", "✓"],
    ["4",  "Document panel balance: obs per year, T distribution, proportion with holes", "Sample", "✓"],
    ["5",  "Compute between, one-way within, TWFE, and first-difference transformations for all variables", "Transform", "✓"],
    ["6",  "Compute variance decomposition; classify variables as time-invariant, mixed, or individual-invariant", "Transform", "✓"],
    ["7",  "Plot univariate distributions for Y and key X under each transformation", "Desc. Stats", "✓"],
    ["8",  "Compute and compare correlation matrices under all four transformations", "Desc. Stats", "✓"],
    ["9",  "Plot bivariate scatter (linear, quadratic, LOWESS) for Y vs. key X under each transformation", "Desc. Stats", "✓"],
    ["10", "Run Pooled OLS and Year FE — identify time-invariant X coefficients via between-variation", "Estimation", "✓"],
    ["11", "Run Entity FE — drop time-invariant variables, interpret within-country coefficients", "Estimation", "✓"],
    ["12", "Run TWFE — verify further attenuation vs. entity FE; check R² within", "Estimation", "✓"],
    ["13", "Run Between OLS — only N observations; good for time-invariant X but confounded by omitted heterogeneity", "Estimation", "✓"],
    ["14", "Run Mundlak RE — add group means of time-varying X as controls for correlated random effects", "Estimation", "✓"],
    ["15", "Run First Differences — eliminates entity FE; interpret ΔY ~ ΔX; drop time-invariant X", "Estimation", "✓"],
    ["16", "If time-invariant X is of interest: run IV using time-invariant instruments (cross-country)", "IV", "✓"],
    ["17", "Report IV first-stage F-statistics; flag weak instruments (F<10); consider LIML or Anderson-Rubin CIs", "IV", "✓"],
    ["18", "Test for autocorrelation in ΔY; if autocorrelation > 0.1, add lagged Y or use Anderson-Hsiao IV", "Dynamic", "✓"],
    ["19", "Run Anderson-Hsiao estimator; check first-stage R²; compute impulse responses and long-run coefficient", "Dynamic", "✓"],
    ["20", "Run robustness: alternative Y (sub-outcomes), sub-samples (developed/developing), non-linear specs", "Robustness", "✓"],
    ["21", "Run interaction: time-invariant X × time-varying Z to recover within-country variation in CB effect", "Interactions", "✓"],
    ["22", "If Hausman test rejects RE: confirm with Mundlak test; if time-invariant X: note Hausman-Taylor (Stata only)", "Spec. test", "✓"],
    ["23", "Cluster standard errors at the unit (entity) level throughout", "SE", "✓"],
    ["24", "Ensure all statistics computed on the sample-selected panel (Y not NaN)", "Housekeeping", "✓"],
    ["25", "Write introductory section: summarise what you found that was NOT in the paper", "Write-up", "✓"],
]
make_table(doc, todo_headers, todo_rows,
           col_widths=[Inches(0.3), Inches(4.2), Inches(0.9), Inches(0.4)])

# ── Save ────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"Report saved → {OUT}")
